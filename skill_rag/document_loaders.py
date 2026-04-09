"""Unified document loaders for multiple file formats."""

import json
import re
import yaml
from abc import ABC, abstractmethod
from pathlib import Path
from dataclasses import dataclass
from typing import List, Optional, Dict, Any, Callable
import logging

from skill_rag.semantic_chunking import create_chunker, SemanticChunker

logger = logging.getLogger(__name__)


def _sanitize_metadata_value(value: Any) -> Any:
    """Sanitize metadata value for ChromaDB compatibility.
    
    ChromaDB only supports: str, int, float, bool, and lists of those types.
    Complex types (dicts, nested lists) are serialized to JSON strings.
    """
    if value is None:
        return ""
    if isinstance(value, (str, int, float, bool)):
        return value
    if isinstance(value, list):
        if all(isinstance(item, (str, int, float, bool)) for item in value):
            return value
        try:
            simple_list = []
            for item in value:
                if isinstance(item, (str, int, float, bool)):
                    simple_list.append(item)
                else:
                    return json.dumps(value, ensure_ascii=False)
            return simple_list
        except (TypeError, ValueError):
            return json.dumps(value, ensure_ascii=False)
    try:
        return json.dumps(value, ensure_ascii=False)
    except (TypeError, ValueError):
        return str(value)


def sanitize_metadata(metadata: Dict[str, Any]) -> Dict[str, Any]:
    """Sanitize all metadata values for ChromaDB compatibility."""
    return {key: _sanitize_metadata_value(value) for key, value in metadata.items()}


@dataclass
class Document:
    """Represents a document chunk with metadata."""
    content: str
    metadata: Dict[str, Any]
    source: str
    chunk_index: int = 0
    start_line: int = 0
    end_line: int = 0


class BaseLoader(ABC):
    """Base class for all document loaders."""
    
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 100,
        use_semantic_chunking: bool = False
    ):
        """Initialize loader with chunking parameters.
        
        Args:
            chunk_size: Maximum characters per chunk
            chunk_overlap: Overlap between chunks
            use_semantic_chunking: Use semantic chunking (paragraph/sentence boundaries)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.use_semantic_chunking = use_semantic_chunking
        self._semantic_chunker: Optional[SemanticChunker] = None
        if use_semantic_chunking:
            self._semantic_chunker = create_chunker(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap
            )
    
    @abstractmethod
    def load_file(self, file_path: Path) -> List[Document]:
        """Load a single file and return document chunks.
        
        Args:
            file_path: Path to file
            
        Returns:
            List of Document chunks
        """
        pass
    
    def load_directory(
        self,
        directory: Path,
        pattern: Optional[str] = None,
        recursive: bool = True
    ) -> List[Document]:
        """Load all matching files from a directory.
        
        Args:
            directory: Directory to search
            pattern: Glob pattern for files (e.g., "*.md", "*.py")
            recursive: Whether to search recursively
            
        Returns:
            List of all Document chunks
        """
        documents = []
        
        if recursive:
            file_paths = list(directory.rglob(pattern or self._get_default_pattern()))
        else:
            file_paths = list(directory.glob(pattern or self._get_default_pattern()))
        
        for file_path in sorted(file_paths):
            try:
                docs = self.load_file(file_path)
                documents.extend(docs)
            except Exception as e:
                logger.warning(f"Failed to load {file_path}: {e}")
        
        return documents
    
    @abstractmethod
    def _get_default_pattern(self) -> str:
        """Get default file pattern for this loader."""
        pass
    
    def _compute_fingerprint(self, content: str) -> str:
        """Compute content fingerprint for deduplication.
        
        Args:
            content: Document content
            
        Returns:
            SHA256 hash of content
        """
        import hashlib
        return hashlib.sha256(content.encode('utf-8')).hexdigest()


class MarkdownLoader(BaseLoader):
    """Load and chunk markdown files with YAML frontmatter support."""
    
    def _get_default_pattern(self) -> str:
        return "*.md"
    
    def load_file(self, file_path: Path) -> List[Document]:
        """Load a single markdown file and return document chunks."""
        content = file_path.read_text(encoding='utf-8')
        frontmatter, body, frontmatter_line_count = self._parse_frontmatter(content)
        
        # Chunk by headers to preserve semantic structure
        chunks = self._chunk_by_headers(body, frontmatter_line_count)
        
        documents = []
        for i, (chunk, start_line, end_line) in enumerate(chunks):
            # Build raw metadata
            raw_metadata = {
                'source': str(file_path),
                'filename': file_path.name,
                'name': frontmatter.get('name', file_path.stem),
                'content_fingerprint': self._compute_fingerprint(chunk),
                **{k: v for k, v in frontmatter.items() if k != 'name'}
            }
            # Sanitize metadata for ChromaDB compatibility
            doc = Document(
                content=chunk,
                metadata=sanitize_metadata(raw_metadata),
                source=str(file_path),
                chunk_index=i,
                start_line=start_line,
                end_line=end_line
            )
            documents.append(doc)
        
        return documents
    
    def _parse_frontmatter(self, content: str) -> tuple[Dict[str, Any], str, int]:
        """Parse YAML frontmatter from markdown content."""
        pattern = r'^---\s*\n(.*?)\n---\s*\n(.*)$'
        match = re.match(pattern, content, re.DOTALL)
        
        if match:
            try:
                frontmatter = yaml.safe_load(match.group(1)) or {}
                body = match.group(2)
                text_before_body = content[:match.start(2)]
                frontmatter_lines = text_before_body.count('\n')
                return frontmatter, body, frontmatter_lines
            except yaml.YAMLError:
                pass
        
        return {}, content, 0
    
    def _chunk_by_headers(self, content: str, frontmatter_offset: int = 0) -> List[tuple[str, int, int]]:
        """Split content by headers (##) to preserve semantic structure."""
        lines = content.splitlines()
        
        section_boundaries = [0]
        for i, line in enumerate(lines):
            if re.match(r'^##\s', line):
                section_boundaries.append(i)
        
        section_boundaries.append(len(lines))
        
        chunks = []
        for i in range(len(section_boundaries) - 1):
            start_idx = section_boundaries[i]
            end_idx = section_boundaries[i + 1]
            
            section_lines = lines[start_idx:end_idx]
            section = '\n'.join(section_lines).strip()
            
            if not section:
                continue
            
            start_line = start_idx + frontmatter_offset + 1
            end_line = end_idx + frontmatter_offset
            
            if len(section) > self.chunk_size:
                chunks.extend(self._split_large_section(section, start_line))
            else:
                chunks.append((section, start_line, end_line))
        
        return chunks if chunks else [(content.strip(), frontmatter_offset + 1, frontmatter_offset + len(lines))]
    
    def _split_large_section(self, section: str, section_start_line: int) -> List[tuple[str, int, int]]:
        """Split a large section into smaller chunks."""
        lines = section.splitlines()
        paragraphs = []
        para_start_lines = []
        
        current_para = []
        current_para_start = 0
        
        for i, line in enumerate(lines):
            if line.strip() == '' and current_para:
                paragraphs.append('\n'.join(current_para).strip())
                para_start_lines.append(section_start_line + current_para_start)
                current_para = []
                current_para_start = i + 1
            else:
                current_para.append(line)
        
        if current_para:
            paragraphs.append('\n'.join(current_para).strip())
            para_start_lines.append(section_start_line + current_para_start)
        
        if not paragraphs:
            return [(section.strip(), section_start_line, section_start_line + len(lines) - 1)]
        
        chunks = []
        current_chunk_paras = []
        current_size = 0
        
        def calc_end_line(para_idx: int, para_text: str) -> int:
            para_start = para_start_lines[para_idx]
            para_line_count = para_text.count('\n') + 1
            return para_start + para_line_count - 1
        
        def save_chunk():
            if current_chunk_paras:
                chunk_text = '\n\n'.join(p[1] for p in current_chunk_paras)
                first_para_idx = current_chunk_paras[0][0]
                last_para_idx = current_chunk_paras[-1][0]
                start = para_start_lines[first_para_idx]
                end = calc_end_line(last_para_idx, current_chunk_paras[-1][1])
                chunks.append((chunk_text, start, end))
        
        for i, para in enumerate(paragraphs):
            if not para:
                continue
            
            para_size = len(para)
            
            if current_size + para_size > self.chunk_size and current_chunk_paras:
                save_chunk()
                
                overlap_count = min(2, len(current_chunk_paras))
                overlap_paras = current_chunk_paras[-overlap_count:]
                current_chunk_paras = overlap_paras + [(i, para)]
                current_size = sum(len(p[1]) for p in current_chunk_paras)
            else:
                current_chunk_paras.append((i, para))
                current_size += para_size
        
        save_chunk()
        
        return chunks


class TextLoader(BaseLoader):
    """Load plain text files and code files."""
    
    # Supported extensions for code files
    CODE_EXTENSIONS = {
        '.py', '.js', '.ts', '.jsx', '.tsx', '.java', '.cpp', '.c', '.h', '.hpp',
        '.cs', '.go', '.rs', '.rb', '.php', '.swift', '.kt', '.scala', '.r',
        '.sql', '.sh', '.bash', '.ps1', '.yaml', '.yml', '.json', '.xml',
        '.html', '.css', '.scss', '.sass', '.less', '.md', '.markdown', '.rst', '.txt',
        '.csv'
    }
    
    def _get_default_pattern(self) -> str:
        return "*"
    
    def load_file(self, file_path: Path) -> List[Document]:
        """Load a text or code file."""
        # Check if it's a supported text file
        if not self._is_text_file(file_path):
            logger.warning(f"Skipping unsupported file: {file_path}")
            return []
        
        try:
            content = file_path.read_text(encoding='utf-8')
        except UnicodeDecodeError:
            logger.warning(f"Cannot decode file as text: {file_path}")
            return []
        
        # Detect language for code files
        language = self._detect_language(file_path)
        
        # Chunk the content
        chunks = self._chunk_text(content)
        
        documents = []
        for i, (chunk, start_line, end_line) in enumerate(chunks):
            metadata = {
                'source': str(file_path),
                'filename': file_path.name,
                'name': file_path.stem,
                'file_type': file_path.suffix.lstrip('.'),
                'language': language,
                'content_fingerprint': self._compute_fingerprint(chunk),
            }
            
            doc = Document(
                content=chunk,
                metadata=sanitize_metadata(metadata),
                source=str(file_path),
                chunk_index=i,
                start_line=start_line,
                end_line=end_line
            )
            documents.append(doc)
        
        return documents
    
    def _is_text_file(self, file_path: Path) -> bool:
        """Check if file is a supported text/code file.
        
        Accepts:
        - Files with known code/text extensions
        - Files without extensions (treated as text files)
        - Files with unknown extensions (attempted as text)
        """
        suffix = file_path.suffix.lower()
        if suffix in self.CODE_EXTENSIONS:
            return True
        # Allow files without extension or with unknown extensions
        # (will be validated during actual read)
        return True
    
    def _detect_language(self, file_path: Path) -> Optional[str]:
        """Detect programming language from file extension."""
        ext_map = {
            '.py': 'python',
            '.js': 'javascript',
            '.ts': 'typescript',
            '.jsx': 'jsx',
            '.tsx': 'tsx',
            '.java': 'java',
            '.cpp': 'cpp',
            '.c': 'c',
            '.h': 'c',
            '.hpp': 'cpp',
            '.cs': 'csharp',
            '.go': 'go',
            '.rs': 'rust',
            '.rb': 'ruby',
            '.php': 'php',
            '.swift': 'swift',
            '.kt': 'kotlin',
            '.scala': 'scala',
            '.r': 'r',
            '.sql': 'sql',
            '.sh': 'bash',
            '.bash': 'bash',
            '.ps1': 'powershell',
            '.yaml': 'yaml',
            '.yml': 'yaml',
            '.json': 'json',
            '.xml': 'xml',
            '.html': 'html',
            '.css': 'css',
            '.scss': 'scss',
            '.sass': 'sass',
            '.less': 'less',
            '.md': 'markdown',
            '.markdown': 'markdown',
            '.rst': 'rst',
            '.txt': 'text',
        }
        return ext_map.get(file_path.suffix.lower())
    
    def _chunk_text(self, content: str) -> List[tuple[str, int, int]]:
        """Split text into chunks while preserving line numbers."""
        lines = content.splitlines()
        total_lines = len(lines)
        
        chunks = []
        current_chunk_lines = []
        current_size = 0
        chunk_start_line = 1
        
        for i, line in enumerate(lines, start=1):
            line_size = len(line) + 1  # +1 for newline
            
            if current_size + line_size > self.chunk_size and current_chunk_lines:
                # Save current chunk
                chunk_text = '\n'.join(current_chunk_lines)
                chunks.append((chunk_text, chunk_start_line, i - 1))
                
                # Start new chunk with overlap
                overlap_lines = current_chunk_lines[-self.chunk_overlap:] if self.chunk_overlap > 0 else []
                current_chunk_lines = overlap_lines + [line]
                current_size = sum(len(l) + 1 for l in current_chunk_lines)
                chunk_start_line = i - len(overlap_lines)
            else:
                current_chunk_lines.append(line)
                current_size += line_size
        
        # Add remaining content
        if current_chunk_lines:
            chunk_text = '\n'.join(current_chunk_lines)
            chunks.append((chunk_text, chunk_start_line, total_lines))
        
        return chunks if chunks else [(content.strip(), 1, max(1, total_lines))]


class PDFLoader(BaseLoader):
    """Load PDF files with optional semantic chunking."""
    
    def _get_default_pattern(self) -> str:
        return "*.pdf"
    
    def load_file(self, file_path: Path) -> List[Document]:
        """Load a PDF file."""
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
            return []
        
        try:
            reader = PdfReader(str(file_path))
            documents = []
            
            full_text = ""
            page_map = []  # Track (start_pos, end_pos, page_num) for each page
            
            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text()
                if text:
                    start_pos = len(full_text)
                    full_text += text + "\n\n"
                    end_pos = len(full_text)
                    page_map.append((start_pos, end_pos, page_num))
            
            if not full_text.strip():
                logger.warning(f"No text extracted from PDF: {file_path}")
                return []
            
            # Use semantic chunking if enabled, otherwise use standard chunking
            if self.use_semantic_chunking and self._semantic_chunker:
                chunks = self._chunk_text_semantic(full_text, page_map)
            else:
                chunks = self._chunk_text_with_page_info(full_text, page_map)
            
            for i, (chunk, start_page, end_page) in enumerate(chunks):
                metadata = {
                    'source': str(file_path),
                    'filename': file_path.name,
                    'name': file_path.stem,
                    'file_type': 'pdf',
                    'start_page': start_page,
                    'end_page': end_page,
                    'total_pages': len(reader.pages),
                    'content_fingerprint': self._compute_fingerprint(chunk),
                }
                
                doc = Document(
                    content=chunk,
                    metadata=sanitize_metadata(metadata),
                    source=str(file_path),
                    chunk_index=i,
                    start_line=start_page,  # Use page number as line reference
                    end_line=end_page
                )
                documents.append(doc)
            
            return documents
            
        except Exception as e:
            logger.error(f"Failed to load PDF {file_path}: {e}")
            return []
    
    def _chunk_text_semantic(
        self,
        text: str,
        page_map: List[tuple[int, int, int]]
    ) -> List[tuple[str, int, int]]:
        """Chunk text using semantic boundaries while tracking pages."""
        chunks = self._semantic_chunker.chunk(text)
        
        result = []
        for chunk in chunks:
            # Find page numbers for this chunk
            start_page = 1
            end_page = 1
            
            for start, end, page in page_map:
                if start <= chunk.start_pos < end:
                    start_page = page
                if start < chunk.end_pos <= end:
                    end_page = page
                    break
            
            result.append((chunk.content, start_page, end_page))
        
        return result if result else [(text.strip(), 1, page_map[-1][2] if page_map else 1)]
    
    def _chunk_text_with_page_info(
        self,
        text: str,
        page_map: List[tuple[int, int, int]]
    ) -> List[tuple[str, int, int]]:
        """Split text into chunks while tracking page numbers."""
        # Simple paragraph-based chunking
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        chunks = []
        current_chunk_paras = []
        current_size = 0
        chunk_start_page = page_map[0][2] if page_map else 1
        chunk_end_page = chunk_start_page
        current_pos = 0
        
        for para in paragraphs:
            para_size = len(para)
            
            # Find page number for this paragraph
            para_page = 1
            for start, end, page in page_map:
                if start <= current_pos < end:
                    para_page = page
                    break
            
            if current_size + para_size > self.chunk_size and current_chunk_paras:
                # Save current chunk
                chunk_text = '\n\n'.join(current_chunk_paras)
                chunks.append((chunk_text, chunk_start_page, chunk_end_page))
                
                # Start new chunk
                current_chunk_paras = [para]
                current_size = para_size
                chunk_start_page = para_page
                chunk_end_page = para_page
            else:
                current_chunk_paras.append(para)
                current_size += para_size
                chunk_end_page = para_page
            
            current_pos += para_size + 2  # +2 for '\n\n'
        
        # Add remaining content
        if current_chunk_paras:
            chunk_text = '\n\n'.join(current_chunk_paras)
            chunks.append((chunk_text, chunk_start_page, chunk_end_page))
        
        return chunks if chunks else [(text.strip(), 1, page_map[-1][2] if page_map else 1)]


class WordLoader(BaseLoader):
    """Load Microsoft Word documents with optional semantic chunking."""
    
    def _get_default_pattern(self) -> str:
        return "*.docx"
    
    def load_file(self, file_path: Path) -> List[Document]:
        """Load a Word document."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return []
        
        try:
            doc = DocxDocument(str(file_path))
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            full_text = '\n\n'.join(paragraphs)
            
            if not full_text.strip():
                logger.warning(f"No text extracted from Word document: {file_path}")
                return []
            
            # Extract document properties
            core_props = doc.core_properties
            base_metadata = {
                'source': str(file_path),
                'filename': file_path.name,
                'name': file_path.stem,
                'file_type': 'docx',
                'title': core_props.title or '',
                'author': core_props.author or '',
                'created': str(core_props.created) if core_props.created else '',
            }
            
            # Use semantic chunking if enabled
            if self.use_semantic_chunking and self._semantic_chunker:
                chunks = self._chunk_text_semantic(full_text)
            else:
                chunks = self._chunk_text(full_text)
            
            documents = []
            for i, (chunk, start_line, end_line) in enumerate(chunks):
                chunk_metadata = {
                    **base_metadata,
                    'chunk_index': i,
                    'content_fingerprint': self._compute_fingerprint(chunk),
                }
                
                doc_obj = Document(
                    content=chunk,
                    metadata=sanitize_metadata(chunk_metadata),
                    source=str(file_path),
                    chunk_index=i,
                    start_line=start_line,
                    end_line=end_line
                )
                documents.append(doc_obj)
            
            return documents
            
        except Exception as e:
            logger.error(f"Failed to load Word document {file_path}: {e}")
            return []
    
    def _chunk_text_semantic(self, content: str) -> List[tuple[str, int, int]]:
        """Chunk text using semantic boundaries."""
        chunks = self._semantic_chunker.chunk(content)
        
        result = []
        for chunk in chunks:
            # Map character positions to paragraph numbers
            start_para = content[:chunk.start_pos].count('\n\n') + 1
            end_para = content[:chunk.end_pos].count('\n\n') + 1
            result.append((chunk.content, start_para, end_para))
        
        return result if result else [(content.strip(), 1, 1)]
    
    def _chunk_text(self, content: str) -> List[tuple[str, int, int]]:
        """Split text into chunks."""
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        
        chunks = []
        current_chunk_paras = []
        current_size = 0
        para_count = 0
        chunk_start_para = 0
        
        for i, para in enumerate(paragraphs):
            para_size = len(para)
            
            if current_size + para_size > self.chunk_size and current_chunk_paras:
                # Save current chunk
                chunk_text = '\n\n'.join(current_chunk_paras)
                chunks.append((chunk_text, chunk_start_para + 1, i))
                
                # Start new chunk
                current_chunk_paras = [para]
                current_size = para_size
                chunk_start_para = i
            else:
                current_chunk_paras.append(para)
                current_size += para_size
        
        # Add remaining content
        if current_chunk_paras:
            chunk_text = '\n\n'.join(current_chunk_paras)
            chunks.append((chunk_text, chunk_start_para + 1, len(paragraphs)))
        
        return chunks if chunks else [(content.strip(), 1, 1)]


class AutoLoader(BaseLoader):
    """Auto-detect file type and use appropriate loader.
    
    Supports semantic chunking for code files (Python, JavaScript, Java, etc.)
    when use_semantic_code_chunking is enabled.
    
    Also supports semantic chunking for general documents (PDF, Word, etc.)
    when use_semantic_chunking is enabled.
    """
    
    # Code file extensions that support semantic chunking
    CODE_EXTENSIONS = {
        '.py', '.pyw', '.js', '.jsx', '.ts', '.tsx', '.mjs',
        '.java', '.cpp', '.cc', '.cxx', '.c', '.h', '.hpp',
        '.go', '.rs', '.cs', '.kt', '.scala'
    }
    
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 100,
        use_semantic_code_chunking: bool = True,
        use_semantic_chunking: bool = False
    ):
        """Initialize auto-loader.
        
        Args:
            chunk_size: Maximum characters per chunk
            chunk_overlap: Overlap between chunks
            use_semantic_code_chunking: Use CodeLoader for semantic code parsing
            use_semantic_chunking: Use semantic chunking for general documents
        """
        super().__init__(chunk_size, chunk_overlap, use_semantic_chunking)
        self.use_semantic_code_chunking = use_semantic_code_chunking
        self._loaders: Dict[str, BaseLoader] = {}
        self._init_loaders()
    
    def _init_loaders(self):
        """Initialize loaders for different file types."""
        from .code_loader import CodeLoader, create_code_loader
        
        # Initialize loaders with semantic chunking setting
        self._loaders = {
            '.md': MarkdownLoader(self.chunk_size, self.chunk_overlap),
            '.txt': TextLoader(self.chunk_size, self.chunk_overlap),
            '.pdf': PDFLoader(
                self.chunk_size,
                self.chunk_overlap,
                use_semantic_chunking=self.use_semantic_chunking
            ),
            '.docx': WordLoader(
                self.chunk_size,
                self.chunk_overlap,
                use_semantic_chunking=self.use_semantic_chunking
            ),
        }
        
        # Use CodeLoader for semantic chunking of code files
        if self.use_semantic_code_chunking:
            try:
                code_loader = create_code_loader(
                    chunk_size=self.chunk_size,
                    chunk_overlap=self.chunk_overlap,
                    include_imports=True
                )
                for ext in self.CODE_EXTENSIONS:
                    self._loaders[ext] = code_loader
            except Exception as e:
                logger.warning(f"Failed to initialize CodeLoader: {e}")
                self._init_text_loaders_for_code()
        else:
            self._init_text_loaders_for_code()
    
    def _init_text_loaders_for_code(self):
        """Use TextLoader for code files (fallback)."""
        for ext in self.CODE_EXTENSIONS:
            self._loaders[ext] = TextLoader(self.chunk_size, self.chunk_overlap)
    
    def _get_default_pattern(self) -> str:
        return "*"
    
    def load_file(self, file_path: Path) -> List[Document]:
        """Load a file using appropriate loader based on extension."""
        file_path = Path(file_path)
        ext = file_path.suffix.lower()
        
        loader = self._loaders.get(ext)
        if loader:
            return loader.load_file(file_path)
        else:
            # Try text loader as fallback
            logger.debug(f"Unknown extension {ext}, trying text loader for {file_path}")
            text_loader = TextLoader(self.chunk_size, self.chunk_overlap)
            return text_loader.load_file(file_path)
    
    def register_loader(self, extension: str, loader: BaseLoader):
        """Register a custom loader for a file extension.
        
        Args:
            extension: File extension (e.g., '.csv')
            loader: Loader instance
        """
        self._loaders[extension.lower()] = loader


def get_loader_for_file(file_path: Path) -> BaseLoader:
    """Get appropriate loader for a file based on extension.
    
    Args:
        file_path: Path to file
        
    Returns:
        Appropriate loader instance
    """
    ext = Path(file_path).suffix.lower()
    
    loaders = {
        '.md': MarkdownLoader,
        '.pdf': PDFLoader,
        '.docx': WordLoader,
    }
    
    loader_class = loaders.get(ext, TextLoader)
    return loader_class()
