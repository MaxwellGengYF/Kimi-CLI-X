"""Document loader for markdown files with frontmatter parsing and multi-format support."""

import io
import json
import re
import yaml
from pathlib import Path
from dataclasses import dataclass
from typing import List, Optional, Dict, Any


def _sanitize_metadata_value(value: Any) -> Any:
    """Sanitize metadata value for ChromaDB compatibility.
    
    ChromaDB only supports: str, int, float, bool, and lists of those types.
    Complex types (dicts, nested lists) are serialized to JSON strings.
    
    Args:
        value: The metadata value to sanitize
        
    Returns:
        Sanitized value compatible with ChromaDB
    """
    if value is None:
        return ""
    if isinstance(value, (str, int, float, bool)):
        return value
    if isinstance(value, list):
        # Check if all items are simple types
        if all(isinstance(item, (str, int, float, bool)) for item in value):
            return value
        # Check if it's a list of strings (after converting non-strings)
        try:
            simple_list = []
            for item in value:
                if isinstance(item, (str, int, float, bool)):
                    simple_list.append(item)
                else:
                    # Contains complex items, serialize the whole list
                    return json.dumps(value, ensure_ascii=False)
            return simple_list
        except (TypeError, ValueError):
            return json.dumps(value, ensure_ascii=False)
    # For dicts or any other complex types, serialize to JSON
    try:
        return json.dumps(value, ensure_ascii=False)
    except (TypeError, ValueError):
        return str(value)


def sanitize_metadata(metadata: Dict[str, Any]) -> Dict[str, Any]:
    """Sanitize all metadata values for ChromaDB compatibility.
    
    Args:
        metadata: Raw metadata dict
        
    Returns:
        Sanitized metadata dict
    """
    return {key: _sanitize_metadata_value(value) for key, value in metadata.items()}


@dataclass
class Document:
    """Represents a document chunk with metadata."""
    content: str
    metadata: Dict[str, Any]
    source: str
    chunk_index: int = 0
    start_line: int = 0    # Starting line number in source file
    end_line: int = 0      # Ending line number in source file


class MarkdownLoader:
    """Load and chunk markdown files with YAML frontmatter support."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 100):
        """Initialize loader with chunking parameters.
        
        Args:
            chunk_size: Maximum characters per chunk
            chunk_overlap: Overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def load_file(self, file_path: Path) -> List[Document]:
        """Load a single markdown file and return document chunks.
        
        Args:
            file_path: Path to markdown file
            
        Returns:
            List of Document chunks
        """
        content = file_path.read_text(encoding='utf-8')
        frontmatter, body, frontmatter_line_count = self._parse_frontmatter(content)
        
        # Chunk by headers to preserve semantic structure
        chunks = self._chunk_by_headers(body, frontmatter_line_count)
        
        documents = []
        for i, (chunk, start_line, end_line) in enumerate(chunks):
            # Build raw metadata
            raw_metadata = {
                'source': str(file_path),
                'filename': file_path.name,
                'name': frontmatter.get('name', file_path.stem),
                **{k: v for k, v in frontmatter.items() if k != 'name'}
            }
            # Sanitize metadata for ChromaDB compatibility
            doc = Document(
                content=chunk,
                metadata=sanitize_metadata(raw_metadata),
                source=str(file_path),
                chunk_index=i,
                start_line=start_line,
                end_line=end_line
            )
            documents.append(doc)
        
        return documents
    
    def load_directory(self, directory: Path, pattern: str = "*.md") -> List[Document]:
        """Load all markdown files from a directory.
        
        Args:
            directory: Directory to search
            pattern: Glob pattern for files
            
        Returns:
            List of all Document chunks
        """
        documents = []
        for file_path in sorted(directory.rglob(pattern)):
            docs = self.load_file(file_path)
            documents.extend(docs)
        return documents
    
    def load_text(
        self,
        content: str,
        source: str = "inline",
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """Load and chunk raw text content.
        
        Args:
            content: Text content
            source: Source identifier
            metadata: Optional metadata
            
        Returns:
            List of Document chunks
        """
        frontmatter, body, frontmatter_line_count = self._parse_frontmatter(content)
        
        # Merge provided metadata with frontmatter
        merged_metadata = {**(metadata or {}), **frontmatter}
        
        # Chunk by headers
        chunks = self._chunk_by_headers(body, frontmatter_line_count)
        
        documents = []
        for i, (chunk, start_line, end_line) in enumerate(chunks):
            raw_metadata = {
                'source': source,
                'filename': source,
                'name': merged_metadata.get('name', source),
                **{k: v for k, v in merged_metadata.items() if k != 'name'}
            }
            doc = Document(
                content=chunk,
                metadata=sanitize_metadata(raw_metadata),
                source=source,
                chunk_index=i,
                start_line=start_line,
                end_line=end_line
            )
            documents.append(doc)
        
        return documents
    
    def _parse_frontmatter(self, content: str) -> tuple[Dict[str, Any], str, int]:
        """Parse YAML frontmatter from markdown content.
        
        Args:
            content: Raw markdown content
            
        Returns:
            Tuple of (frontmatter dict, body content, frontmatter line count)
        """
        # Match YAML frontmatter between --- markers
        pattern = r'^---\s*\n(.*?)\n---\s*\n(.*)$'
        match = re.match(pattern, content, re.DOTALL)
        
        if match:
            try:
                frontmatter = yaml.safe_load(match.group(1)) or {}
                body = match.group(2)
                # Count lines before body starts (more accurate than counting in group(1))
                text_before_body = content[:match.start(2)]
                frontmatter_lines = text_before_body.count('\n')
                return frontmatter, body, frontmatter_lines
            except yaml.YAMLError:
                pass
        
        # No valid frontmatter found
        return {}, content, 0
    
    def _chunk_by_headers(self, content: str, frontmatter_offset: int = 0) -> List[tuple[str, int, int]]:
        """Split content by headers (##) to preserve semantic structure.
        
        Args:
            content: Markdown body content
            frontmatter_offset: Number of lines to add to line numbers (from frontmatter)
            
        Returns:
            List of tuples: (chunk_content, start_line, end_line)
        """
        lines = content.splitlines()
        
        # Split by ## headers while tracking line numbers
        section_boundaries = [0]  # Start of first section
        for i, line in enumerate(lines):
            if re.match(r'^##\s', line):
                section_boundaries.append(i)
        
        # Add end boundary
        section_boundaries.append(len(lines))
        
        chunks = []
        for i in range(len(section_boundaries) - 1):
            start_idx = section_boundaries[i]
            end_idx = section_boundaries[i + 1]
            
            # Extract section lines
            section_lines = lines[start_idx:end_idx]
            section = '\n'.join(section_lines).strip()
            
            if not section:
                continue
            
            # Calculate line numbers (1-indexed, including frontmatter offset)
            start_line = start_idx + frontmatter_offset + 1  # +1 for 1-indexed
            end_line = end_idx + frontmatter_offset  # Don't add +1 here, end_idx is exclusive
            
            # If section is too long, split further by paragraphs
            if len(section) > self.chunk_size:
                chunks.extend(self._split_large_section(section, start_line))
            else:
                chunks.append((section, start_line, end_line))
        
        return chunks if chunks else [(content.strip(), frontmatter_offset + 1, frontmatter_offset + len(lines))]
    
    def _split_large_section(self, section: str, section_start_line: int) -> List[tuple[str, int, int]]:
        """Split a large section into smaller chunks while preserving line number info.
        
        Args:
            section: Large text section
            section_start_line: Starting line number of this section in the source file
            
        Returns:
            List of tuples: (chunk_content, start_line, end_line)
        """
        lines = section.splitlines()
        paragraphs = []
        para_start_lines = []
        
        current_para = []
        current_para_start = 0
        
        for i, line in enumerate(lines):
            if line.strip() == '' and current_para:
                # End of paragraph
                paragraphs.append('\n'.join(current_para).strip())
                para_start_lines.append(section_start_line + current_para_start)
                current_para = []
                current_para_start = i + 1
            else:
                current_para.append(line)
        
        # Add last paragraph
        if current_para:
            paragraphs.append('\n'.join(current_para).strip())
            para_start_lines.append(section_start_line + current_para_start)
        
        if not paragraphs:
            return [(section.strip(), section_start_line, section_start_line + len(lines) - 1)]
        
        chunks = []
        current_chunk_paras = []  # List of (para_index, para_text)
        current_chunk_start_line = para_start_lines[0]
        current_size = 0
        
        def calc_end_line(para_idx: int, para_text: str) -> int:
            """Calculate end line for a paragraph."""
            para_start = para_start_lines[para_idx]
            para_lines = para_text.count('\n') + 1
            return para_start + para_lines - 1
        
        def save_chunk():
            """Save current chunk to results."""
            if current_chunk_paras:
                chunk_text = '\n\n'.join(p[1] for p in current_chunk_paras)
                first_para_idx = current_chunk_paras[0][0]
                last_para_idx = current_chunk_paras[-1][0]
                start = para_start_lines[first_para_idx]
                end = calc_end_line(last_para_idx, current_chunk_paras[-1][1])
                chunks.append((chunk_text, start, end))
        
        for i, para in enumerate(paragraphs):
            if not para:
                continue
                
            para_size = len(para)
            
            if current_size + para_size > self.chunk_size and current_chunk_paras:
                # Save current chunk
                save_chunk()
                
                # Start new chunk with overlap (last 1-2 paragraphs)
                overlap_count = min(2, len(current_chunk_paras))
                overlap_paras = current_chunk_paras[-overlap_count:]
                current_chunk_paras = overlap_paras + [(i, para)]
                current_chunk_start_line = para_start_lines[overlap_paras[0][0]]
                current_size = sum(len(p[1]) for p in current_chunk_paras)
            else:
                current_chunk_paras.append((i, para))
                current_size += para_size
        
        # Add remaining content
        save_chunk()
        
        return chunks


class PDFLoader:
    """Load text from PDF files."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 100):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def load_file(self, file_path: Path) -> List[Document]:
        """Load text from a PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            List of Document chunks
        """
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF loading. Install with: pip install PyPDF2")
        
        reader = PdfReader(str(file_path))
        
        documents = []
        text_buffer = []
        current_size = 0
        chunk_index = 0
        current_page = 0
        
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if not text.strip():
                continue
            
            current_page = page_num + 1
            text_buffer.append(text)
            current_size += len(text)
            
            # When buffer exceeds chunk size, create chunks
            if current_size >= self.chunk_size:
                full_text = '\n'.join(text_buffer)
                chunks = self._chunk_text(full_text, current_page - len(text_buffer) + 1)
                
                for chunk_text, start_page in chunks:
                    doc = Document(
                        content=chunk_text,
                        metadata={
                            'source': str(file_path),
                            'filename': file_path.name,
                            'name': file_path.stem,
                            'page_start': start_page,
                            'page_end': current_page,
                            'format': 'pdf'
                        },
                        source=str(file_path),
                        chunk_index=chunk_index,
                        start_line=0,  # PDFs don't have lines
                        end_line=0
                    )
                    documents.append(doc)
                    chunk_index += 1
                
                text_buffer = []
                current_size = 0
        
        # Process remaining text
        if text_buffer:
            full_text = '\n'.join(text_buffer)
            chunks = self._chunk_text(full_text, current_page - len(text_buffer) + 1)
            
            for chunk_text, start_page in chunks:
                doc = Document(
                    content=chunk_text,
                    metadata={
                        'source': str(file_path),
                        'filename': file_path.name,
                        'name': file_path.stem,
                        'page_start': start_page,
                        'page_end': current_page,
                        'format': 'pdf'
                    },
                    source=str(file_path),
                    chunk_index=chunk_index,
                    start_line=0,
                    end_line=0
                )
                documents.append(doc)
                chunk_index += 1
        
        return documents
    
    def _chunk_text(self, text: str, start_page: int) -> List[tuple[str, int]]:
        """Split text into chunks while tracking page numbers.
        
        Returns:
            List of (chunk_text, start_page) tuples
        """
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        if not paragraphs:
            return [(text.strip(), start_page)] if text.strip() else []
        
        chunks = []
        current_chunk = []
        current_size = 0
        chunk_start_page = start_page
        
        for para in paragraphs:
            para_size = len(para)
            
            if current_size + para_size > self.chunk_size and current_chunk:
                # Save current chunk
                chunks.append(('\n\n'.join(current_chunk), chunk_start_page))
                
                # Start new chunk with overlap
                overlap_count = min(2, len(current_chunk))
                current_chunk = current_chunk[-overlap_count:] + [para]
                current_size = sum(len(p) for p in current_chunk)
            else:
                current_chunk.append(para)
                current_size += para_size
        
        # Add remaining content
        if current_chunk:
            chunks.append(('\n\n'.join(current_chunk), chunk_start_page))
        
        return chunks


class DocxLoader:
    """Load text from Word (.docx) files."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 100):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def load_file(self, file_path: Path) -> List[Document]:
        """Load text from a Word file.
        
        Args:
            file_path: Path to .docx file
            
        Returns:
            List of Document chunks
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx is required for Word loading. Install with: pip install python-docx")
        
        doc = DocxDocument(str(file_path))
        
        # Extract text from paragraphs
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)
        
        documents = []
        current_chunk = []
        current_size = 0
        chunk_index = 0
        
        for para in paragraphs:
            para_size = len(para)
            
            if current_size + para_size > self.chunk_size and current_chunk:
                # Save current chunk
                doc = Document(
                    content='\n\n'.join(current_chunk),
                    metadata={
                        'source': str(file_path),
                        'filename': file_path.name,
                        'name': file_path.stem,
                        'format': 'docx'
                    },
                    source=str(file_path),
                    chunk_index=chunk_index,
                    start_line=0,
                    end_line=0
                )
                documents.append(doc)
                chunk_index += 1
                
                # Start new chunk with overlap
                overlap_count = min(2, len(current_chunk))
                current_chunk = current_chunk[-overlap_count:] + [para]
                current_size = sum(len(p) for p in current_chunk)
            else:
                current_chunk.append(para)
                current_size += para_size
        
        # Add remaining content
        if current_chunk:
            doc = Document(
                content='\n\n'.join(current_chunk),
                metadata={
                    'source': str(file_path),
                    'filename': file_path.name,
                    'name': file_path.stem,
                    'format': 'docx'
                },
                source=str(file_path),
                chunk_index=chunk_index,
                start_line=0,
                end_line=0
            )
            documents.append(doc)
        
        return documents


class UniversalDocumentLoader:
    """Universal document loader that supports multiple file formats.
    
    Supports semantic chunking for all document types when use_semantic_chunking is True.
    """
    
    SUPPORTED_FORMATS = {
        '.md': MarkdownLoader,
        '.markdown': MarkdownLoader,
        '.pdf': PDFLoader,
        '.docx': DocxLoader,
    }
    
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 100,
        use_semantic_chunking: bool = False
    ):
        """Initialize universal loader.
        
        Args:
            chunk_size: Maximum characters per chunk
            chunk_overlap: Overlap between chunks
            use_semantic_chunking: Use semantic chunking with paragraph/sentence boundaries
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.use_semantic_chunking = use_semantic_chunking
        self._loaders: Dict[str, Any] = {}
        self._semantic_chunker: Optional[Any] = None
        
        if use_semantic_chunking:
            try:
                from skill_rag.semantic_chunking import create_chunker
                self._semantic_chunker = create_chunker(
                    chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap
                )
            except Exception as e:
                import logging
                logging.getLogger(__name__).warning(
                    f"Failed to initialize semantic chunker: {e}. Falling back to standard chunking."
                )
                self.use_semantic_chunking = False
    
    def _get_loader(self, extension: str) -> Any:
        """Get or create loader for a file extension."""
        if extension not in self._loaders:
            loader_class = self.SUPPORTED_FORMATS.get(extension)
            if loader_class:
                self._loaders[extension] = loader_class(
                    chunk_size=self.chunk_size,
                    chunk_overlap=self.chunk_overlap
                )
        return self._loaders.get(extension)
    
    def load_file(self, file_path: Path) -> List[Document]:
        """Load a file based on its extension.
        
        Args:
            file_path: Path to file
            
        Returns:
            List of Document chunks
            
        Raises:
            ValueError: If file format is not supported
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        loader = self._get_loader(extension)
        if loader:
            return loader.load_file(file_path)
        
        raise ValueError(f"Unsupported file format: {extension}. "
                        f"Supported: {list(self.SUPPORTED_FORMATS.keys())}")
    
    def load_directory(
        self,
        directory: Path,
        pattern: str = "*"
    ) -> List[Document]:
        """Load all supported files from a directory.
        
        Args:
            directory: Directory to search
            pattern: File pattern (e.g., "*.md" or "*" for all supported)
            
        Returns:
            List of all Document chunks
        """
        directory = Path(directory)
        documents = []
        
        if pattern == "*":
            # Load all supported formats
            for ext in self.SUPPORTED_FORMATS.keys():
                for file_path in sorted(directory.rglob(f"*{ext}")):
                    try:
                        docs = self.load_file(file_path)
                        documents.extend(docs)
                    except Exception as e:
                        print(f"Warning: Failed to load {file_path}: {e}")
        else:
            # Load specific pattern
            for file_path in sorted(directory.rglob(pattern)):
                try:
                    docs = self.load_file(file_path)
                    documents.extend(docs)
                except ValueError as e:
                    # Skip unsupported formats
                    pass
                except Exception as e:
                    print(f"Warning: Failed to load {file_path}: {e}")
        
        return documents
    
    def load_text(
        self,
        content: str,
        source: str = "inline",
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """Load and chunk raw text content (treated as markdown).
        
        Args:
            content: Text content
            source: Source identifier
            metadata: Optional metadata
            
        Returns:
            List of Document chunks
        """
        loader = MarkdownLoader(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap
        )
        return loader.load_text(content, source, metadata)
    
    @classmethod
    def is_supported(cls, file_path: Path) -> bool:
        """Check if a file format is supported.
        
        Args:
            file_path: Path to file
            
        Returns:
            True if supported
        """
        extension = Path(file_path).suffix.lower()
        return extension in cls.SUPPORTED_FORMATS
