from kimi_agent_sdk import CallableTool2, ToolError, ToolOk, ToolReturnValue
from pydantic import BaseModel, Field
import re
import os
from pathlib import Path
from my_tools.common import _maybe_export_output

class Params(BaseModel):
    docx_path: str = Field(
        description="Path to the DOCX file to convert.",
    )
    output_path: str = Field(
        default="",
        description="Path for the output markdown file (optional). If not provided, returns markdown content directly.",
    )


class DocxToMarkdown(CallableTool2):
    name: str = "DocxToMarkdown"
    description: str = "Convert Word documents (DOCX) to Markdown format."
    params: type[Params] = Params

    async def __call__(self, params: Params) -> ToolReturnValue:
        # Validate DOCX path
        if not os.path.exists(params.docx_path):
            return ToolError(
                output="",
                message=f"DOCX file not found: {params.docx_path}",
                brief="File not found",
            )

        # Check file extension
        ext = os.path.splitext(params.docx_path)[1].lower()
        if ext != ".docx":
            return ToolError(
                output="",
                message=f"Invalid file format: {ext}. Only .docx files are supported.",
                brief="Invalid file format",
            )

        # Try importing required libraries
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return ToolError(
                output="",
                message="python-docx is not installed. Install with: pip install python-docx",
                brief="Missing dependency: python-docx",
            )

        try:
            # Perform conversion
            converter = _DocxToMarkdownConverter(params.docx_path)
            markdown_content = converter.convert()

            if params.output_path:
                # Save to file
                output_dir = os.path.dirname(os.path.abspath(params.output_path))
                if output_dir and not os.path.exists(output_dir):
                    os.makedirs(output_dir, exist_ok=True)
                
                Path(params.output_path).write_text(markdown_content, encoding='utf-8')
                
                return ToolOk(
                    output=f"Markdown saved to: {params.output_path}",
                    message=f"Converted {len(markdown_content)} characters.",
                )
            else:
                # Return content directly (truncated if too long)
                max_length = 10000
                if len(markdown_content) > max_length:
                    truncated = markdown_content[:max_length] + "\n\n... [content truncated]"
                    return ToolOk(
                        output=_maybe_export_output(truncated, params.docx_path),
                        message=f"Content was truncated (total: {len(markdown_content)} characters). Use output_path parameter to save full content to file.",
                    )
                return ToolOk(
                    output=_maybe_export_output(markdown_content, params.docx_path),
                    message=f"Converted {len(markdown_content)} characters.",
                )

        except Exception as e:
            return ToolError(
                output="",
                message=f"DOCX conversion failed: {str(e)}",
                brief="Conversion failed",
            )


class _DocxToMarkdownConverter:
    """Internal converter class for DOCX to Markdown conversion."""
    
    def __init__(self, docx_path):
        from docx import Document
        self.doc = Document(docx_path)
        self.markdown_lines = []
        
    def convert(self):
        """执行转换"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        for para in self.doc.paragraphs:
            md_line = self._convert_paragraph(para)
            if md_line:
                self.markdown_lines.append(md_line)
        
        # 处理表格
        for table in self.doc.tables:
            self.markdown_lines.append(self._convert_table(table))
            
        return '\n'.join(self.markdown_lines)
    
    def _convert_paragraph(self, para):
        """转换单个段落"""
        text = para.text.strip()
        if not text:
            return ''
        
        # 获取样式名称
        style_name = para.style.name if para.style else 'Normal'
        
        # 判断对齐方式
        alignment = ''
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            alignment = 'center'
        elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            alignment = 'right'
        
        # 处理标题
        if style_name.startswith('Heading'):
            level = self._extract_heading_level(style_name)
            return f"{'#' * level} {text}"
        
        # 处理标题样式（如 Title, Subtitle）
        if style_name == 'Title':
            return f"# {text}"
        if style_name == 'Subtitle':
            return f"## {text}"
        
        # 处理列表
        if style_name.startswith('List Bullet'):
            return f"- {text}"
        if style_name.startswith('List Number'):
            return f"1. {text}"
        
        # 处理代码块（通过样式名判断）
        if 'Code' in style_name or 'Source' in style_name:
            return f"```\n{text}\n```"
        
        # 处理引用
        if 'Quote' in style_name:
            return f"> {text}"
        
        # 普通段落 - 处理行内格式
        return self._process_inline_formatting(para)
    
    def _extract_heading_level(self, style_name):
        """提取标题级别"""
        match = re.search(r'Heading\s*(\d)', style_name)
        return int(match.group(1)) if match else 1
    
    def _process_inline_formatting(self, para):
        """处理行内格式（粗体、斜体、代码等）"""
        result = []
        
        for run in para.runs:
            text = run.text
            if not text:
                continue
                
            # 处理格式
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            
            # 处理代码样式（通过字体判断）
            if run.font and run.font.name in ['Courier New', 'Consolas', 'Monaco']:
                text = f"`{text}`"
            
            # 处理删除线
            if run.font and run.font.strike:
                text = f"~~{text}~~"
            
            result.append(text)
        
        return ''.join(result)
    
    def _convert_table(self, table):
        """转换表格为 Markdown 表格"""
        if not table.rows:
            return ''
        
        lines = []
        
        # 表头
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        lines.append('| ' + ' | '.join(header_cells) + ' |')
        
        # 分隔符
        separators = ['---'] * len(header_cells)
        lines.append('| ' + ' | '.join(separators) + ' |')
        
        # 数据行
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append('| ' + ' | '.join(cells) + ' |')
        
        return '\n'.join(lines)
